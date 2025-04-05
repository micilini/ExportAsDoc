import json
import base64
import sys  # Import sys for command-line arguments
from io import BytesIO
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from PIL import Image
from docx.shared import RGBColor  # Import RGBColor
from docx.enum.text import WD_COLOR_INDEX  # Import color index for highlighting
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor
from docx.oxml import parse_xml
from docx.shared import Inches


def customize_run(run, font_size=None, color=None):
    """
    Customize a run with the given font size and color.
    :param run: The run object to customize.
    :param font_size: Font size in points.
    :param color: Color as an RGB tuple (e.g., (255, 0, 0) for red).
    """
    if font_size:
        run.font.size = Pt(font_size)
    if color:
        run.font.color.rgb = RGBColor(*color)

#when we add any hyper links this added twice one as a hyperlink and the another as normal text, so this function will delete duplicated normal text.
def clean_duplicate_links(doc):
    for paragraph in doc.paragraphs:
        # Set to store text of hyperlinks
        hyperlink_texts = set()
        runs_to_remove = []

        # Iterate through the paragraph's XML to find hyperlink elements
        for element in paragraph._element:
            if element.tag.endswith("hyperlink"):
                # Extract text inside the hyperlink
                for child in element:
                    if child.tag.endswith("r"):  # Check if it's a run element
                        for t in child:
                            if t.tag.endswith("t"):  # Text element
                                hyperlink_text = t.text.strip()
                                hyperlink_texts.add(hyperlink_text)

        # Check runs for duplicates
        for run in paragraph.runs:
            if run.text.strip() in hyperlink_texts and run._element.getparent().tag != qn("w:hyperlink"):
                # Mark duplicate plain-text runs for removal
                runs_to_remove.append(run)

        # Remove duplicate runs
        for run in runs_to_remove:
            paragraph._element.remove(run._element)

#the function which added hyperlink
def add_hyperlink(paragraph, text, url):
    # Create the hyperlink tag
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    # Create a run and set its properties for the hyperlink
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    # Set font properties for the hyperlink text
    run_pr = OxmlElement("w:rPr")
    run_font = OxmlElement("w:sz")
    run_font.set(qn('w:val'), '24')  # Adjust size
    run_pr.append(run_font)
    run_bold = OxmlElement("w:b")
    run_bold.set(qn('w:val'), '1')  # Bold
    run_pr.append(run_bold)
    run_underline = OxmlElement("w:u")
    run_underline.set(qn('w:val'), 'single')  # Underline
    run_pr.append(run_underline)
    run_color = OxmlElement("w:color")
    run_color.set(qn('w:val'), '0000FF')  # Blue color for hyperlink
    run_pr.append(run_color)
    run.append(run_pr)
    text_element = OxmlElement("w:t")
    text_element.text = text
    run.append(text_element)
    hyperlink.append(run)

    # Add the hyperlink to the paragraph
    paragraph._element.append(hyperlink)



# Load JSON file
def load_json(file_path):
    with open(file_path, 'r', encoding='utf-8') as file:
        return json.load(file)

#this one for make quote like textbox
def style_as_textbox(paragraph, background_color="D9D9D9"):
    """
    Styles a paragraph to look like a textbox with borders and background shading.
    :param paragraph: The paragraph to style.
    :param background_color: Background color in hex (default is light gray 'D9D9D9').
    """
    # Get the paragraph's XML
    p = paragraph._element

    # Create a border element
    pPr = p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")

    for border_type in ["top", "left", "bottom", "right"]:
        border = OxmlElement(f"w:{border_type}")
        border.set(qn("w:val"), "single")  # Border type
        border.set(qn("w:sz"), "4")  # Border width
        border.set(qn("w:space"), "1")  # Space between border and text
        border.set(qn("w:color"), "000000")  # Black border color
        borders.append(border)

    pPr.append(borders)

    # Add background shading
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), background_color)  # Background color
    pPr.append(shading)


def process_block(block, doc):
    block_type = block.get('type')
    data = block.get('data', {})

    if block_type == 'header':
        level = data.get('level', 1)
        text = data.get('text', '')
        paragraph = doc.add_heading(level=level)
        add_formatted_text(paragraph, text)
        for run in paragraph.runs:
            if level==1:
                customize_run(run, font_size=32, color=(0, 0, 0))  # Black headers, size decreases with level
            else:
                customize_run(run, font_size=20 + (6 - level) * 2, color=(0, 0, 0))  # Black headers, size decreases with level
        
    elif block_type == 'paragraph':
        paragraph = doc.add_paragraph()
        add_formatted_text(paragraph, data.get('text', ''))
        for run in paragraph.runs:
           run.font.size = Pt(18)

    elif block_type == 'list':
        style = data.get('style', 'unordered')
        for item in data.get('items', []):
            paragraph = doc.add_paragraph(style='List Bullet' if style == 'unordered' else 'List Number')
            add_formatted_text(paragraph, item)
            for run in paragraph.runs:
               run.font.size = Pt(18)

    elif block_type == 'checklist':
        for item in data.get('items', []):
            status = "✅" if item.get('checked') else "⬜"
            paragraph = doc.add_paragraph(f"{status} ")
            add_formatted_text(paragraph, item.get('text', ''))
            for run in paragraph.runs:
              run.font.size = Pt(18)

    elif block_type == 'quote':
        paragraph = doc.add_paragraph(style='Intense Quote')
        add_formatted_text(paragraph, f"“{data.get('text', '')}”")
        style_as_textbox(paragraph, background_color="F4F4F4")  # Light gray for quotes
        for run in paragraph.runs:
           run.font.size = Pt(18)
        caption = data.get('caption', '')
        if caption:
            caption_paragraph = doc.add_paragraph(style='Caption')
            add_formatted_text(caption_paragraph, f"- {caption}")
            for run in paragraph.runs:
              run.font.size = Pt(14)

    elif block_type == 'warning': 
        title = data.get('title', '')
        message = data.get('message', '')
        paragraph = doc.add_paragraph(style='Quote')
        add_formatted_text(paragraph, f"⚠️ {title}: {message}")
        style_as_textbox(paragraph, background_color="FFF2CC")  # Light yellow for warnings
        for run in paragraph.runs:
           run.font.size = Pt(18)

    elif block_type == 'code':
        paragraph = doc.add_paragraph(data.get('code', ''), style='Normal') 
        for run in paragraph.runs:
           run.font.size = Pt(16)

    elif block_type == 'delimiter':
        paragraph = doc.add_paragraph(style='Title')
        paragraph.alignment=WD_PARAGRAPH_ALIGNMENT.CENTER
        add_formatted_text(paragraph, '***')  # Visual separation

    elif block_type == 'table':
        table_data = data.get('content', [])
        table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
        for i, row in enumerate(table_data):
            for j, cell in enumerate(row):
                cell_paragraph = table.cell(i, j).paragraphs[0]
                add_formatted_text(cell_paragraph, cell)
                for run in cell_paragraph.runs:
                  run.font.size = Pt(18)

    elif block_type == 'image':
        add_image_to_fit_page_or_original(data.get('url'), doc)


# Add formatted text with inline styles
def add_formatted_text(paragraph, text):
    from html.parser import HTMLParser

    class InlineStyleParser(HTMLParser):
        def __init__(self):
            super().__init__()
            self.current_tag = None
            self.current_attrs = {}

        def handle_starttag(self, tag, attrs):
            self.current_tag = tag
            self.current_attrs = dict(attrs)

        def handle_endtag(self, tag):
            self.current_tag = None

        def handle_data(self, data):
            run = paragraph.add_run(data)

            if self.current_tag == "i":
               run.italic = True
            elif self.current_tag == "b":
               run.bold = True
            elif self.current_tag == "u":
               run.underline = True
            elif self.current_tag == "code":
               run.font.color.rgb = RGBColor(255, 0, 0)  # Red for code
            elif self.current_tag == "mark":
               run.font.highlight_color = WD_COLOR_INDEX.YELLOW  # Yellow highlight for sublime
            elif self.current_tag == "a":
          # Check for href only if it's a new <a> tag to avoid processing the same link twice
             if "href" in self.current_attrs:
               add_hyperlink(paragraph, data, self.current_attrs["href"])
          # Reset to None after processing
               self.current_tag = None
            else:
             # Reset formatting for normal text
               run.font.color.rgb = None
 
    parser = InlineStyleParser()
    parser.feed(text)


def add_image_to_fit_page_or_original(base64_string, doc, page_width=6.0, page_height=8.0):
    """
    Add an image to the Word document, resizing only if it exceeds the page size.
    
    Args:
        base64_string (str): The base64-encoded image string.
        doc (Document): The Word document object.
        page_width (float): Maximum width of the image in inches (default is 6 inches for standard margins).
        page_height (float): Maximum height of the image in inches (default is 8 inches for standard margins).
    """
    if base64_string.startswith("data:image/"):
        # Decode the base64 image
        img_data = base64.b64decode(base64_string.split(",")[1])
        img = Image.open(BytesIO(img_data))
        
        # Get the original dimensions of the image
        original_width, original_height = img.size  # In pixels
        dpi = img.info.get("dpi", (96, 96))[0]  # Default DPI is 96 if not specified
        img_width_in_inches = original_width / dpi
        img_height_in_inches = original_height / dpi
        
        # Check if the image needs resizing
        if img_width_in_inches > page_width or img_height_in_inches > page_height:
            # Calculate the scaling factors to fit the image within the page
            width_scale = page_width / img_width_in_inches
            height_scale = page_height / img_height_in_inches
            scale = min(width_scale, height_scale)  # Use the smaller scaling factor
            
            # Calculate the new dimensions
            scaled_width = img_width_in_inches * scale
            scaled_height = img_height_in_inches * scale
        else:
            # Use the original dimensions
            scaled_width = img_width_in_inches
            scaled_height = img_height_in_inches
        
        # Save the image temporarily
        img.save("temp_image.png")
        
        # Add visual separation before the image
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        paragraph.add_run('\n')  # Visual separation
        
        # Add the image to the document with calculated dimensions
        doc.add_picture("temp_image.png", width=Inches(scaled_width), height=Inches(scaled_height))
        
        # Add visual separation after the image
        paragraph = doc.add_paragraph()
        paragraph.add_run('\n')  # Visual separation

# Main function
def json_to_docx(json_file, output_file):
    data = load_json(json_file)
    doc = Document()

    for block in data.get('blocks', []):
        process_block(block, doc)

    # Clean duplicate links
    clean_duplicate_links(doc)

    doc.save(output_file)

# for test
def save_docx_stream_to_file(docx_stream, filename="ExportedNote.docx"):
    """
    Save the DOCX binary stream to a file in the local AppData directory.
    """
    local_appdata_path = os.path.join(os.getenv("LOCALAPPDATA"), "MinimalTextEditorLite")
    os.makedirs(local_appdata_path, exist_ok=True)

    file_path = os.path.join(local_appdata_path, filename)

    with open(file_path, 'wb') as f:
        f.write(docx_stream)


# Run the script
# Define a new function to process arguments
def process_arguments():
    """
    Process command-line arguments.
    Expects JSON as an input and outputs the .docx binary to stdout.
    """
    if len(sys.argv) < 2:
        print("Usage: your_executable_name.exe <json-input>")
        sys.exit(1)

    json_input = sys.argv[1]  # The JSON string or path
    try:
        # Parse JSON directly or load from file
        if json_input.strip().startswith('{'):
            data = json.loads(json_input)
        else:
            with open(json_input, 'r', encoding='utf-8') as file:
                data = json.load(file)

        # Create the Word document
        doc = Document()
        for block in data.get('blocks', []):
            process_block(block, doc)
        clean_duplicate_links(doc)
        # Save to a BytesIO object
        output = BytesIO()
        doc.save(output)

        # Output the binary DOCX data to stdout
        sys.stdout.buffer.write(output.getvalue())
        # Save Doc in Temporary File
        #save_docx_stream_to_file(output.getvalue(),"ExportedNote.docx")

    except Exception as e:
        print(f"Error: {e}")
        sys.exit(1)


# Main execution
if __name__ == "__main__":
    process_arguments()